#loops over input folder, matches table format 
import os
from docx import Document
import csv
import xml.etree.ElementTree as ET
from pdf2docx import Converter

## SETUP ##

# Define the input folder path
adan_input = "/Users/adana/Downloads/Parsing/Parsing/input"
ashar_input = "/Users/asharfarooq/Downloads/Uliza/Subtitles/inputcsv"
input_folder_path = ashar_input  # Update with your input folder path

# Define the output folder path
adan_output = "/Users/adana/Downloads/Parsing/Parsing/output"
ashar_output = "/Users/asharfarooq/Downloads/Uliza/Subtitles/outputcsv"
output_folder_path = ashar_output  # Update with your output folder path

# Create the output folder if it doesn't exist
if not os.path.exists(output_folder_path):
    os.makedirs(output_folder_path)

# Iterate over each file in the input folder
for file_name in os.listdir(input_folder_path):
    # Construct the full file path
    file_path = os.path.join(input_folder_path, file_name)

    # Check if the file is a PDF or DOC/DOCX
    if file_name.endswith('.pdf'):
    # Convert PDF to DOCX
        docx_file_path = os.path.join(output_folder_path, f'{os.path.splitext(file_name)[0]}.docx')
        cv = Converter(file_path)
        cv.convert(docx_file_path)
        cv.close()

        # Load the converted DOCX file
        doc = Document(docx_file_path)

    elif file_name.endswith('.doc') or file_name.endswith('.docx'):
        # Load the Word document
        doc = Document(file_path)
    else:
        # Skip the file if it's not a PDF or DOC/DOCX
        continue

    # Create the root element of the XML tree
    root = ET.Element('root')

    # Process the Word document and generate XML tree

    # Define the color to language mapping
    color_to_language = {
        'FF0000': 'nam',
        '000000': 'eng',
        '0000FF': 'por',
        'FF9900': 'deu',
        '38761D': 'nuu',
        '1155CC': 'afr',
        '006FC0': 'afr',
        '00AF50': 'nuu',
        '008000': 'nuu', 
        '0070C0': 'afr', 

        # Add more color to language mappings as needed
    }

## SETUP ##

## XML GENERATION  ##

    # Iterate through each table in the document
    for table_index, table in enumerate(doc.tables):
        # Check if it is the first table [SPECIAL CASE OF FIRST TABLE]
        if table_index == 0:
            # Create a new XML element for the table
            table_element = ET.SubElement(root, 'table')

            # Iterate through each row in the table
            for row_index, row in enumerate(list(table.rows)):
                # Create a new XML element for the row
                row_element = ET.SubElement(table_element, 'row')

                # Iterate through each cell in the row
                for cell_index, cell in enumerate(row.cells):
                    # Create a new XML element for the cell
                    cell_element = ET.SubElement(row_element, 'cell')

                    # Extract the text from the cell
                    cell_text = cell.text

                    # Add the text to the cell element
                    cell_element.text = cell_text

        else:

            # Create a new XML element for the table
            
            table_element = ET.SubElement(root, 'table')
            # Iterate through each row in the table
            for row_index , row in enumerate(list(table.rows)):

                # SPECIAL CASE OF TOP 2 ROWS OF METADATA IN THE SECOND TABLE IN THE DOCUMENT
                if table_index == 1 and row_index in [0,1]:
                    # Create a new XML element for the row
                    row_element = ET.SubElement(table_element, 'row')

                    # Iterate through each cell in the row
                    for cell_index, cell in enumerate(row.cells):
                        # Create a new XML element for the cell
                        cell_element = ET.SubElement(row_element, 'cell')

                        # Extract the text from the cell
                        cell_text = cell.text

                        # Add the text to the cell element
                        cell_element.text = cell_text
                    continue

                # Create a new XML element for the row
                row_element = ET.SubElement(table_element, 'row')

                # Iterate through each cell in the row
                counter = 1
                notes = ""
                for cell in row.cells:
                    # Create a new XML element for the cell
                    cell_element = ET.SubElement(row_element, 'cell')

                    # Acquire the notes/metadata information from the 4th cell of the row
                    if counter == 4:
                        for paragraph in cell.paragraphs:
                            notes += paragraph.text
                        counter = 1

                        # Adding a notes tag
                        notes_element = ET.SubElement(cell_element, "notes")
                        notes_element.text = notes
                        continue
                    else:
                        counter += 1 

                    # Create a list to store language-specific text
                    language_text_list = []
                    # Iterate through each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        # Create a dictionary to store run-specific text and color
                        run_text_color = {}

                        # Iterate through each run in the paragraph
                        for run in paragraph.runs:

                            # Extract the text color from the run
                            text_color = None
                            if run.font.color.rgb is None:
                                text_color = "000000"
                            else:
                                text_color = str(run.font.color.rgb)

                            # Extract the text from the run
                            run_text = run.text

                            # Extracting run text that is the speaker name(identified by ': ' in the conversation) [SPEAKER NOT IN USE CURRENTLY]
                            # if run_text and run_text[-2:] == ": ":
                            #         speaker = run_text[:-2]

                            #         # # Acquiring speaker name when speaker is off screen (OS)
                            #         # if speaker.startswith("OS"):
                            #         #     # Case of OS (name)
                            #         #     if "(" in speaker:
                            #         #         starting_index = speaker.index("(")
                            #         #         ending_index = speaker.index(")")
                            #         #         speaker = speaker[starting_index+1:ending_index]
                            #         #     # Case of OS 2 (no parentheses)
                            #         #     else:
                            #         #         speaker = speaker[3:]
                                
                            #         # Creating speaker tag
                            #         speaker_element = ET.SubElement(cell_element, "speaker")
                            #         speaker_element.text = speaker
                            # Extracting timestamp information
                            if run_text.count(":") == 2:
                                    # Creating time tag
                                    time_element = ET.SubElement(cell_element, "time")
                                    time_element.text = run_text
                            else:
                                # Store the text and color in the dictionary 
                                if text_color in run_text_color:
                                    run_text_color[text_color] += run_text
                                else:
                                    run_text_color[text_color] = run_text

                        # Create language-specific text
                        for color, text in run_text_color.items():
                            language = color_to_language.get(color)
                            if language:
                                language_text_list.append((language, text))
                            else:
                                language_text_list.append(('sub-cell', text))

                    # Create language tags or sub-cell tags with corresponding text
                    for language, text in language_text_list:
                        if language == 'sub-cell':
                            sub_cell_element = ET.SubElement(cell_element, 'sub-cell')
                            sub_cell_element.text = text
                        else:

                            # # If English, also extracting the speaker and adding a speaker tag  (second column speaker) [SPEAKER NOT IN USE CURRENTLY]
                            # if language == 'eng' and ":" in text and not run_text.count(":") == 2:
                            #     speaker = text[:text.index(":")]

                            #     # Acquiring speaker name when speaker is off screen (OS)
                            #     if speaker.startswith("OS"):
                            #         # Case of OS (name)
                            #         if "(" in speaker:
                            #             starting_index = speaker.index("(")
                            #             ending_index = speaker.index(")")
                            #             speaker = speaker[starting_index+1:ending_index]
                            #         # Case of OS 2 (no parentheses)
                            #         else:
                            #             speaker = speaker[3:]

                            #     # Creating speaker tag
                            #     speaker_element = ET.SubElement(cell_element, "speaker")
                            #     speaker_element.text = speaker
                            
                            # Adding the rest of the language tags
                            language_element = ET.SubElement(cell_element, language)
                            language_element.text = text

    # Create an XML tree from the root element
    tree = ET.ElementTree(root)

    # Write the XML tree to a file
    xml_file_path = os.path.join(output_folder_path, f'{os.path.splitext(file_name)[0]}.xml')
    tree.write(xml_file_path)

## XML GENERATION  ##

## CSV GENERATION ##

    # Getting the header rows content
    section = doc.sections[0]
    header = section.header
    header_row_data = []

    if header.tables:
        table = header.tables[0]

        # Access rows and cells in the header table
        for row in table.rows:
            for cell in row.cells:
                # Store the content of each cell in the header table
                header_row_data.append(cell.text)


    # Generate the CSV file path
    csv_file_path = os.path.join(output_folder_path, f'{os.path.splitext(file_name)[0]}.csv')

    # Generate the metadata CSV file path
    metadata_csv_file_path = os.path.join(output_folder_path, f'{os.path.splitext(file_name)[0]} metadata.csv')

    # Open the CSV files in write mode
    with open(csv_file_path, 'w', newline='', encoding='utf-8') as csv_file, open(metadata_csv_file_path, 'w', newline='', encoding='utf-8') as metadata_csv_file:
        writer = csv.writer(csv_file)
        metadata_writer = csv.writer(metadata_csv_file)

        # Adding header row to the CSV file
        writer.writerow(header_row_data)

        # Iterate through each table in the XML
        for table_index, table in enumerate(root.iter('table')):
            # Check if it is the first table [SPECIAL CASE OF FIRST TABLE OF DOCUMENT]
            if table_index == 0:
                # Iterate through each row in the table
                for row in table.iter('row'):
                    # Create a list to store the data for each cell in the row
                    row_data = []

                    # Iterate through each cell in the row
                    for cell in row.iter('cell'):
                        # Append the text of the cell to the row data list
                        row_data.append(cell.text)

                    # Write the row data to the metadata CSV file
                    metadata_writer.writerow(row_data)
            else:
                #skip the odd tables [SPECIAL CASE OF SKIPPING THE REPETITIVE TOP TABLE OF EACH PAGE AFTER FIRST PAGE]
                if table_index%2 == 0:
                    continue 
                else:
                # Create a list to store the data for each row
                    table_data = []

                    # Iterate through each row in the table
                    for row_index , row in enumerate(table.iter('row')):
                        # Create a list to store the data for each cell in the row

                        # SPECIAL CASE OF TOP 2 ROWS OF METADATA IN THE SECOND TABLE IN THE DOCUMENT
                        if row_index in [0,1] and table_index == 1: 
                            row_data = []

                            # Iterate through each cell in the row
                            for cell_index , cell in enumerate(row.iter('cell')):
                                if cell_index > 0:
                                    continue 
                                # Append the text of the cell to the row data list
                                row_data.append(cell.text)

                            # Write the row data to the metadata CSV file
                            metadata_writer.writerow(row_data)
                            continue 

                        row_data = []



                        # Iterate through each cell in the row
                        for cell in row.iter('cell'):
                            # Create a list to store the data for each language or sub-cell in the cell
                            cell_data = []

                            # Iterate through each language or sub-cell tag in the cell
                            for element in cell:
                                # Check if the element is a language tag or sub-cell tag
                                if element.tag == 'sub-cell':
                                    # Handle sub-cell
                                    text = element.text
                                # Not adding any tags for these xml tags in the CSV
                                elif element.tag in ['notes', 'time', 'eng']:
                                    text = element.text
                                else:
                                    # Handle language tag
                                    text = f"<{element.tag}>{element.text}</{element.tag}>"

                                # Append the text to the cell data list
                                cell_data.append(text)

                            # Join the cell data using a separator and append it to the row data list
                            row_data.append(' '.join(cell_data))

                        # Append the row data to the table data list
                        table_data.append(row_data)

                    # Write the table data to the CSV files
                    writer.writerows(table_data)

    print(f"Processed file: {file_name}")

## CSV GENERATION ##

print("Conversion complete.")
